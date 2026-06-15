#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Xuat tai lieu Rui ro trien khai Winsum Home ra .docx."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUTPUT = Path(__file__).resolve().parent / "Ru-ro-trien-khai-Winsum-Home.docx"

PROJECT = "Website Winsum Home — Quản lý cửa hàng bán đèn & nội thất"
DOC_TYPE = "Rủi ro triển khai"

SECTIONS: list[dict] = [
    {
        "title": "1. Rủi ro môi trường cài đặt",
        "rows": [
            ("Phụ thuộc XAMPP", "Apache + MySQL + PHP 8.0+ trên Windows. Máy demo thiếu XAMPP hoặc sai phiên bản PHP có thể không chạy được.", "Trung bình"),
            ("Xung đột cổng", "Port 80 (Apache) hoặc 3306 (MySQL) bị chiếm bởi IIS/Skype → website hoặc DB không khởi động.", "Trung bình"),
            ("Sai đường dẫn / cấu hình", "Copy project sai thư mục htdocs, import nhầm DB, hoặc config/database.php không khớp → lỗi kết nối.", "Thấp"),
        ],
        "note": "Biện pháp: Hướng dẫn chi tiết tại docs/huong-dan-cai-dat.md; URL mặc định http://localhost/webwinsum.",
    },
    {
        "title": "2. Rủi ro dữ liệu và CSDL",
        "rows": [
            ("Import SQL thất bại", "File dump lớn, phpMyAdmin có thể timeout nếu giới hạn upload/php thấp.", "Trung bình"),
            ("Lệch schema", "Cần chạy script migrate bổ sung (docs/sql/). Bỏ qua có thể làm coupon/tồn kho không khớp code.", "Trung bình"),
            ("Mất dữ liệu demo", "Không có backup tự động trên máy local; cài lại XAMPP có thể mất đơn hàng mẫu.", "Thấp"),
        ],
        "note": "",
    },
    {
        "title": "3. Rủi ro bảo mật (nếu đưa lên internet)",
        "rows": [
            ("Tài khoản demo", "admin / admin123 ghi trong README — nguy hiểm nếu public mà không đổi mật khẩu.", "Cao"),
            ("HTTP không mã hóa", "Localhost chưa HTTPS; deploy production cần SSL.", "Cao"),
            ("MySQL root rỗng", "Cấu hình XAMPP mặc định chỉ phù hợp dev.", "Cao"),
        ],
        "note": "Phiên bản hiện tại chỉ phục vụ demo và bảo vệ đồ án trên máy local.",
    },
    {
        "title": "4. Rủi ro nghiệp vụ và thanh toán",
        "rows": [
            ("VietQR demo", "Khách tự bấm «Đã thanh toán»; không có webhook ngân hàng đối soát thật.", "Cao"),
            ("COD thủ công", "Admin phải cập nhật paid sau giao hàng; quên thao tác → doanh thu dashboard sai.", "Trung bình"),
            ("Luồng hoàn 4 bước", "Admin thao tác sai thứ tự có thể lệch tồn kho / doanh thu.", "Trung bình"),
        ],
        "note": "",
    },
    {
        "title": "5. Rủi ro hiệu năng",
        "rows": [
            ("Giỏ hàng Session", "Không đồng bộ đa thiết bị; session hết hạn có thể mất giỏ.", "Thấp (demo)"),
            ("Race condition tồn kho", "Chưa test nhiều người đặt cùng lúc sản phẩm cuối cùng.", "Trung bình"),
            ("Chưa test tải", "Chỉ kiểm thử thủ công trên một máy.", "Trung bình"),
        ],
        "note": "",
    },
]

CONCLUSION = (
    "Hệ thống đủ điều kiện demo và bảo vệ đồ án (161/161 test case Pass). "
    "Triển khai kinh doanh thật cần: HTTPS, đổi credential, cổng thanh toán có webhook, "
    "backup CSDL, khóa tồn kho bằng transaction, và bổ sung test tự động (PHPUnit)."
)

INTRO = (
    "Tài liệu này mô tả các rủi ro nhóm em nhận thấy khi cài đặt, chạy thử và chuẩn bị demo đồ án. "
    "Hệ thống hiện phù hợp môi trường XAMPP local, chưa sẵn sàng vận hành thương mại công khai "
    "nếu không xử lý thêm các rủi ro mức cao."
)

LEVEL_COLORS = {
    "Cao": RGBColor(0xC0, 0x39, 0x2B),
    "Trung bình": RGBColor(0xE6, 0x7E, 0x22),
    "Thấp": RGBColor(0x27, 0xAE, 0x60),
}


def set_font(run, *, bold: bool = False, size: int = 13, color: RGBColor | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def set_cell(cell, bold: bool = False, size: int = 11, color: RGBColor | None = None) -> None:
    for p in cell.paragraphs:
        for r in p.runs:
            set_font(r, bold=bold, size=size, color=color)
        if not p.runs and p.text:
            r = p.add_run(p.text)
            p.text = ""
            set_font(r, bold=bold, size=size, color=color)


def add_risk_table(doc: Document, rows: list[tuple[str, str, str]]) -> None:
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Rủi ro", "Mô tả", "Mức độ"]
    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr[i].text = text
        set_cell(hdr[i], bold=True)
    for risk, desc, level in rows:
        row = table.add_row().cells
        row[0].text = risk
        row[1].text = desc
        row[2].text = level
        for c in (0, 1):
            set_cell(row[c])
        color = None
        for key, rgb in LEVEL_COLORS.items():
            if key in level:
                color = rgb
                break
        set_cell(row[2], bold=("Cao" in level), color=color)
    doc.add_paragraph()


def build_document() -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(13)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(t.add_run(DOC_TYPE), bold=True, size=16)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(sub.add_run(f"Đồ án: {PROJECT}"), bold=True, size=14)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(meta.add_run("Cập nhật: 06/2026 | Nhóm phát triển đồ án"), size=12)

    doc.add_paragraph()
    p = doc.add_paragraph()
    set_font(p.add_run(INTRO))
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for section in SECTIONS:
        h = doc.add_heading(section["title"], level=1)
        for r in h.runs:
            set_font(r, bold=True, size=14)
        add_risk_table(doc, section["rows"])
        if section.get("note"):
            note = doc.add_paragraph()
            set_font(note.add_run(section["note"]), bold=True, size=12)

    h = doc.add_heading("6. Kết luận", level=1)
    for r in h.runs:
        set_font(r, bold=True, size=14)
    concl = doc.add_paragraph()
    set_font(concl.add_run(CONCLUSION))
    concl.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(foot.add_run("Tài liệu sinh từ docs/generate-ru-ro-doc-word.py — Winsum Home"), size=11, color=RGBColor(0x66, 0x66, 0x66))

    return doc


def main() -> None:
    doc = build_document()
    doc.save(OUTPUT)
    print(str(OUTPUT))


if __name__ == "__main__":
    main()
