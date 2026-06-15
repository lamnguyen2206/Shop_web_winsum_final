#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Xuất tài liệu đặc tả CSDL Winsum Home ra .docx
theo mẫu: Lược đồ dữ liệu → Bảng liên kết (K/C) → ERD.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUTPUT = Path(__file__).resolve().parent / "Mo-ta-CSDL-Winsum-Home.docx"

PROJECT_NAME = "Website Winsum Home — Quản lý cửa hàng bán đèn & nội thất"
DOC_TYPE = "Tài liệu đặc tả thiết kế phần mềm"

# (tên_trường, kiểu, độ_rộng, ràng_buộc, mô_tả)
Field = tuple[str, str, str, str, str]

TABLES: list[dict] = [
    {
        "no": 1,
        "name": "customers",
        "purpose": "Lưu tài khoản khách hàng và quản trị viên (role = admin).",
        "fields": [
            ("id", "BIGINT", "—", "PK, AI", "Mã khách hàng"),
            ("customer_code", "VARCHAR", "30", "UK, NOT NULL", "Mã khách do hệ thống sinh"),
            ("full_name", "VARCHAR", "120", "NOT NULL", "Họ và tên"),
            ("phone", "VARCHAR", "30", "UK, NOT NULL", "Số điện thoại đăng nhập"),
            ("email", "VARCHAR", "120", "UK, NULL", "Email"),
            ("password_hash", "VARCHAR", "255", "NULL", "Mật khẩu mã hóa bcrypt"),
            ("status", "ENUM", "—", "DEFAULT active", "active / inactive / blocked"),
            ("role", "ENUM", "—", "NOT NULL", "customer hoặc admin"),
        ],
    },
    {
        "no": 2,
        "name": "categories",
        "purpose": "Phân loại sản phẩm (đèn thả, đèn bàn, kệ trang trí…).",
        "fields": [
            ("id", "BIGINT", "—", "PK, AI", "Mã danh mục"),
            ("parent_id", "BIGINT", "—", "FK, NULL", "Danh mục cha (tự tham chiếu)"),
            ("name", "VARCHAR", "120", "NOT NULL", "Tên danh mục"),
            ("slug", "VARCHAR", "140", "UK, NOT NULL", "Đường dẫn URL"),
            ("is_active", "TINYINT", "1", "DEFAULT 1", "Trạng thái hiển thị"),
        ],
    },
    {
        "no": 3,
        "name": "products",
        "purpose": "Thông tin sản phẩm bán trên website.",
        "fields": [
            ("id", "BIGINT", "—", "PK, AI", "Mã sản phẩm"),
            ("category_id", "BIGINT", "—", "FK, NOT NULL", "Danh mục sản phẩm"),
            ("name", "VARCHAR", "255", "NOT NULL", "Tên sản phẩm"),
            ("slug", "VARCHAR", "300", "UK, NOT NULL", "Slug URL"),
            ("sku", "VARCHAR", "60", "UK, NOT NULL", "Mã SKU"),
            ("base_price", "DECIMAL", "12,2", "NOT NULL", "Giá bán (VNĐ)"),
            ("stock_status", "ENUM", "—", "DEFAULT in_stock", "in_stock / out_of_stock"),
            ("rating_average", "DECIMAL", "3,2", "DEFAULT 0", "Điểm đánh giá trung bình"),
            ("is_active", "TINYINT", "1", "DEFAULT 1", "Hiển thị trên web"),
        ],
    },
    {
        "no": 4,
        "name": "product_images",
        "purpose": "Lưu nhiều ảnh cho một sản phẩm.",
        "fields": [
            ("id", "BIGINT", "—", "PK, AI", "Mã ảnh"),
            ("product_id", "BIGINT", "—", "FK, NOT NULL", "Sản phẩm liên kết"),
            ("image_url", "VARCHAR", "255", "NOT NULL", "Đường dẫn file ảnh"),
            ("is_primary", "TINYINT", "1", "DEFAULT 0", "Ảnh đại diện"),
        ],
    },
    {
        "no": 5,
        "name": "inventory_items",
        "purpose": "Số lượng tồn theo sản phẩm (1 SP = 1 dòng tồn).",
        "fields": [
            ("id", "BIGINT", "—", "PK, AI", "Mã bản ghi tồn"),
            ("product_id", "BIGINT", "—", "FK, UK, NOT NULL", "Sản phẩm"),
            ("quantity_on_hand", "INT", "—", "NOT NULL", "Số lượng tồn thực tế"),
            ("reorder_level", "INT", "—", "DEFAULT 0", "Ngưỡng cảnh báo nhập thêm"),
        ],
    },
    {
        "no": 6,
        "name": "coupons",
        "purpose": "Quản lý mã giảm giá, freeship.",
        "fields": [
            ("id", "BIGINT", "—", "PK, AI", "Mã coupon"),
            ("code", "VARCHAR", "50", "UK, NOT NULL", "Mã nhập (VD: WINSUMXINCHAO)"),
            ("name", "VARCHAR", "120", "NOT NULL", "Tên chương trình"),
            ("discount_type", "ENUM", "—", "NOT NULL", "fixed / percent / shipping"),
            ("discount_value", "DECIMAL", "12,2", "NOT NULL", "Giá trị giảm"),
            ("per_customer_limit", "INT", "—", "NULL", "Giới hạn lượt dùng/khách"),
            ("is_active", "TINYINT", "1", "DEFAULT 1", "Kích hoạt"),
        ],
    },
    {
        "no": 7,
        "name": "orders",
        "purpose": "Header đơn hàng; lưu snapshot thông tin khách tại thời điểm đặt.",
        "fields": [
            ("id", "BIGINT", "—", "PK, AI", "Mã đơn nội bộ"),
            ("order_code", "VARCHAR", "30", "UK, NOT NULL", "Mã đơn hiển thị (WS…)"),
            ("customer_id", "BIGINT", "—", "FK, NULL", "Khách (NULL nếu guest)"),
            ("customer_name", "VARCHAR", "120", "NOT NULL", "Tên người nhận"),
            ("customer_phone", "VARCHAR", "30", "NOT NULL", "Số điện thoại"),
            ("customer_address", "TEXT", "—", "NOT NULL", "Địa chỉ giao hàng"),
            ("coupon_id", "BIGINT", "—", "FK, NULL", "Mã giảm giá áp dụng"),
            ("subtotal", "DECIMAL", "12,2", "NOT NULL", "Tổng tiền hàng"),
            ("shipping_fee", "DECIMAL", "12,2", "NOT NULL", "Phí vận chuyển"),
            ("discount_amount", "DECIMAL", "12,2", "NOT NULL", "Số tiền giảm"),
            ("grand_total", "DECIMAL", "12,2", "NOT NULL", "Tổng thanh toán"),
            ("payment_status", "ENUM", "—", "NOT NULL", "unpaid / paid / refunded"),
            ("fulfillment_status", "ENUM", "—", "NOT NULL", "Trạng thái giao hàng"),
            ("status", "VARCHAR", "30", "NOT NULL", "Trạng thái đơn tổng hợp"),
        ],
    },
    {
        "no": 8,
        "name": "order_items",
        "purpose": "Thực thể yếu — chi tiết từng dòng sản phẩm trong đơn.",
        "fields": [
            ("id", "BIGINT", "—", "PK, AI", "Mã dòng đơn"),
            ("order_id", "BIGINT", "—", "FK, NOT NULL", "Đơn hàng"),
            ("product_id", "BIGINT", "—", "FK, NOT NULL", "Sản phẩm"),
            ("product_sku", "VARCHAR", "60", "NOT NULL", "SKU tại thời điểm mua"),
            ("product_name", "VARCHAR", "255", "NOT NULL", "Tên SP tại thời điểm mua"),
            ("unit_price", "DECIMAL", "12,2", "NOT NULL", "Đơn giá"),
            ("quantity", "INT", "—", "NOT NULL", "Số lượng đặt"),
            ("line_total", "DECIMAL", "12,2", "NOT NULL", "Thành tiền dòng"),
        ],
    },
    {
        "no": 9,
        "name": "order_return_requests",
        "purpose": "Khách gửi yêu cầu hoàn hàng trong 7 ngày sau giao.",
        "fields": [
            ("id", "BIGINT", "—", "PK, AI", "Mã yêu cầu"),
            ("order_id", "BIGINT", "—", "FK, NOT NULL", "Đơn cần hoàn"),
            ("customer_id", "BIGINT", "—", "FK, NOT NULL", "Khách gửi yêu cầu"),
            ("reason", "VARCHAR", "80", "NOT NULL", "Lý do hoàn"),
            ("evidence_url", "VARCHAR", "255", "NOT NULL", "Ảnh bằng chứng"),
            ("bank_account_number", "VARCHAR", "30", "NOT NULL", "Số tài khoản ngân hàng"),
            ("status", "ENUM", "—", "NOT NULL", "pending / accepted / completed…"),
        ],
    },
    {
        "no": 10,
        "name": "payment_methods",
        "purpose": "Danh mục phương thức thanh toán (COD, VietQR).",
        "fields": [
            ("id", "BIGINT", "—", "PK, AI", "Mã phương thức"),
            ("code", "VARCHAR", "30", "UK, NOT NULL", "cod / bank_transfer"),
            ("name", "VARCHAR", "120", "NOT NULL", "Tên hiển thị"),
        ],
    },
    {
        "no": 11,
        "name": "order_payments",
        "purpose": "Giao dịch thanh toán gắn với đơn hàng.",
        "fields": [
            ("id", "BIGINT", "—", "PK, AI", "Mã giao dịch"),
            ("order_id", "BIGINT", "—", "FK, NOT NULL", "Đơn hàng"),
            ("payment_method_id", "BIGINT", "—", "FK", "Phương thức thanh toán"),
            ("amount", "DECIMAL", "12,2", "NOT NULL", "Số tiền thanh toán"),
            ("status", "ENUM", "—", "NOT NULL", "pending / success / refunded"),
        ],
    },
    {
        "no": 12,
        "name": "shipping_methods",
        "purpose": "Danh mục phương thức vận chuyển và phí ship.",
        "fields": [
            ("id", "BIGINT", "—", "PK, AI", "Mã phương thức"),
            ("code", "VARCHAR", "30", "UK, NOT NULL", "express_24h / standard"),
            ("name", "VARCHAR", "120", "NOT NULL", "Tên hiển thị"),
            ("fee", "DECIMAL", "12,2", "NOT NULL", "Phí ship cố định"),
        ],
    },
    {
        "no": 13,
        "name": "order_shipments",
        "purpose": "Thông tin giao hàng của từng đơn.",
        "fields": [
            ("id", "BIGINT", "—", "PK, AI", "Mã vận chuyển"),
            ("order_id", "BIGINT", "—", "FK, NOT NULL", "Đơn hàng"),
            ("shipping_method_id", "BIGINT", "—", "FK", "Phương thức ship"),
            ("recipient_address", "TEXT", "—", "NOT NULL", "Địa chỉ nhận hàng"),
            ("shipping_fee", "DECIMAL", "12,2", "NOT NULL", "Phí vận chuyển"),
        ],
    },
    {
        "no": 14,
        "name": "product_reviews",
        "purpose": "Khách đánh giá sau khi đơn delivered.",
        "fields": [
            ("id", "BIGINT", "—", "PK, AI", "Mã đánh giá"),
            ("product_id", "BIGINT", "—", "FK, NOT NULL", "Sản phẩm"),
            ("customer_id", "BIGINT", "—", "FK, NULL", "Khách đánh giá"),
            ("rating", "TINYINT", "—", "NOT NULL", "Điểm 1–5"),
            ("content", "TEXT", "—", "NULL", "Nội dung đánh giá"),
        ],
    },
    {
        "no": 15,
        "name": "blog_categories",
        "purpose": "Phân loại bài viết blog.",
        "fields": [
            ("id", "BIGINT", "—", "PK, AI", "Mã danh mục"),
            ("name", "VARCHAR", "120", "UK, NOT NULL", "Tên danh mục"),
            ("slug", "VARCHAR", "140", "UK, NOT NULL", "Slug URL"),
        ],
    },
    {
        "no": 16,
        "name": "blog_posts",
        "purpose": "Nội dung bài viết tin tức / hướng dẫn.",
        "fields": [
            ("id", "BIGINT", "—", "PK, AI", "Mã bài viết"),
            ("category_id", "BIGINT", "—", "FK", "Danh mục blog"),
            ("slug", "VARCHAR", "255", "UK, NOT NULL", "Slug URL"),
            ("title", "VARCHAR", "255", "NOT NULL", "Tiêu đề"),
            ("content", "LONGTEXT", "—", "NOT NULL", "Nội dung HTML"),
            ("published_at", "DATE", "—", "NOT NULL", "Ngày đăng"),
        ],
    },
    {
        "no": 17,
        "name": "blog_comments",
        "purpose": "Bình luận bài viết; admin duyệt trước khi hiển thị.",
        "fields": [
            ("id", "BIGINT", "—", "PK, AI", "Mã bình luận"),
            ("post_id", "BIGINT", "—", "FK, NOT NULL", "Bài viết"),
            ("author_name", "VARCHAR", "120", "NOT NULL", "Tên người bình luận"),
            ("content", "TEXT", "—", "NOT NULL", "Nội dung"),
        ],
    },
]

# (table_name, fields) — id = PK; *_id = FK
SCHEMA_LINES: list[tuple[str, list[str]]] = [
    ("customers", ["id", "customer_code", "full_name", "phone", "email", "password_hash", "status", "role"]),
    ("categories", ["id", "parent_id", "name", "slug", "is_active"]),
    ("products", ["id", "category_id", "name", "slug", "sku", "base_price", "stock_status", "is_active"]),
    ("product_images", ["id", "product_id", "image_url", "is_primary"]),
    ("inventory_items", ["id", "product_id", "quantity_on_hand", "reorder_level"]),
    ("coupons", ["id", "code", "name", "discount_type", "discount_value", "is_active"]),
    (
        "orders",
        [
            "id", "order_code", "customer_id", "coupon_id", "customer_name", "customer_phone",
            "customer_address", "subtotal", "shipping_fee", "discount_amount", "grand_total",
            "payment_status", "status",
        ],
    ),
    (
        "order_items",
        ["id", "order_id", "product_id", "product_sku", "product_name", "unit_price", "quantity", "line_total"],
    ),
    (
        "order_return_requests",
        ["id", "order_id", "customer_id", "reason", "evidence_url", "bank_account_number", "status"],
    ),
    ("payment_methods", ["id", "code", "name"]),
    ("order_payments", ["id", "order_id", "payment_method_id", "amount", "status"]),
    ("shipping_methods", ["id", "code", "name", "fee"]),
    ("order_shipments", ["id", "order_id", "shipping_method_id", "recipient_address", "shipping_fee"]),
    ("product_reviews", ["id", "product_id", "customer_id", "rating", "content"]),
    ("blog_categories", ["id", "name", "slug"]),
    ("blog_posts", ["id", "category_id", "slug", "title", "content", "published_at"]),
    ("blog_comments", ["id", "post_id", "author_name", "content"]),
]

FK_FIELDS = {
    "parent_id", "category_id", "product_id", "customer_id", "coupon_id", "order_id",
    "payment_method_id", "shipping_method_id", "post_id",
}

# 7 bảng cốt lõi — map số thứ tự → tên bảng trong CSDL
LINK_TABLE_MAP = [
    "(1) customers",
    "(2) products",
    "(3) categories",
    "(4) orders",
    "(5) order_items",
    "(6) coupons",
    "(7) order_return_requests",
]

# Mỗi dòng: [thuộc_tính_khóa, c1..c7, liên_kết]
# K = khóa chính tại bảng đó; C = khóa ngoại tham chiếu bảng đó
LINK_MATRIX = [
    ["customer_id", "K", "", "", "C", "", "", "C", "(1)-(4), (1)-(7)"],
    ["product_id", "", "K", "", "", "C", "", "", "(2)-(5)"],
    ["category_id", "", "C", "K", "", "", "", "", "(3)-(2)"],
    ["order_id", "", "", "", "K", "C", "", "C", "(4)-(5), (4)-(7)"],
    ["coupon_id", "", "", "", "C", "", "K", "", "(6)-(4)"],
    ["order_id, product_id", "", "", "", "", "", "", "", "dòng loại"],
    ["order_id, customer_id", "", "", "", "", "", "", "", "dòng loại"],
]

ERD_TEXT = r"""
                    ┌─────────────────────┐
                    │  (3) categories     │
                    │  # id               │
                    │    name, slug       │
                    └──────────┬──────────┘
                               │ 1
                               │ thuộc
                               ▼ N
┌──────────────────┐    ┌─────────────────────┐    ┌──────────────────┐
│ (1) customers    │    │  (2) products       │    │ (6) coupons      │
│ # id             │    │  # id               │    │ # id             │
│   full_name      │    │    category_id (FK) │    │   code, name     │
│   phone          │    │    name, sku, price │    └────────┬─────────┘
└────────┬─────────┘    └──────────┬──────────┘             │
         │ 1                       │ 1                        │ 0..1
         │ đặt                     │ có trong                 │ áp dụng
         ▼ N                       ▼ N                        ▼
┌────────────────────────────────────────────────────────────────────┐
│                     (4) orders                                     │
│  # id, order_code                                                  │
│    customer_id (FK), coupon_id (FK)                                │
│    customer_name, phone, address, grand_total, status              │
└───────────────┬───────────────────────────────┬────────────────────┘
                │ 1                             │ 1
                │ gồm                           │ có thể
                ▼ N                             ▼ 0..N
     ┌─────────────────────┐         ┌──────────────────────────┐
     │ (5) order_items     │         │ (7) order_return_requests│
     │ # id                │         │ # id                     │
     │   order_id (FK)     │         │   order_id (FK)          │
     │   product_id (FK)   │         │   customer_id (FK)       │
     │   quantity, price   │         │   reason, status         │
     └─────────────────────┘         └──────────────────────────┘

Ghi chú: # = Khóa chính (PK), gạch chân trong lược đồ = Khóa ngoại (FK).
Giỏ hàng lưu PHP Session (không có bảng carts).
"""


def set_run_font(run, bold: bool = False, size: int = 13, color: RGBColor | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def set_cell_font(cell, bold: bool = False, size: int = 11) -> None:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            set_run_font(run, bold=bold, size=size)
        if not paragraph.runs and paragraph.text:
            run = paragraph.add_run(paragraph.text)
            paragraph.text = ""
            set_run_font(run, bold=bold, size=size)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        set_run_font(run, bold=True, size=14 if level == 1 else 13)


def add_paragraph(doc: Document, text: str, justify: bool = True, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, bold=bold)
    if justify:
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_schema_table(doc: Document, fields: list[Field]) -> None:
    headers = ["STT", "Tên trường", "Kiểu dữ liệu", "Độ rộng", "Ràng buộc", "Mô tả"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr[i].text = text
        set_cell_font(hdr[i], bold=True)
    for idx, field in enumerate(fields, start=1):
        row = table.add_row().cells
        row[0].text = str(idx)
        for c, val in enumerate(field, start=1):
            row[c].text = val
        for cell in row:
            set_cell_font(cell)
    doc.add_paragraph()


def add_link_matrix(doc: Document) -> None:
    headers = ["Thuộc tính khóa", "(1)", "(2)", "(3)", "(4)", "(5)", "(6)", "(7)", "Liên kết"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr[i].text = text
        set_cell_font(hdr[i], bold=True)
    for row_data in LINK_MATRIX:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = val
            set_cell_font(row[i], bold=(val in ("K", "C")))
    doc.add_paragraph()


def add_mono_run(paragraph, text: str, *, bold: bool = False, italic: bool = False, underline: bool = False) -> None:
    run = paragraph.add_run(text)
    run.font.name = "Courier New"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
    run.font.size = Pt(9)
    run.bold = bold
    run.italic = italic
    run.underline = underline


def add_schema_block(doc: Document) -> None:
    for table_name, fields in SCHEMA_LINES:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        add_mono_run(p, f"{table_name} (", bold=True)
        for i, field in enumerate(fields):
            if i > 0:
                add_mono_run(p, ", ")
            if field == "id":
                add_mono_run(p, field, underline=True)
            elif field in FK_FIELDS:
                add_mono_run(p, field, italic=True)
            else:
                add_mono_run(p, field)
        add_mono_run(p, ")")
    doc.add_paragraph()


def add_mono_block(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text.strip())
    run.font.name = "Courier New"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
    run.font.size = Pt(9)


def build_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(13)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    # Trang bìa
    t1 = doc.add_paragraph()
    t1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t1.add_run(DOC_TYPE)
    set_run_font(r, bold=True, size=16)

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t2.add_run(f"Dự án: {PROJECT_NAME}")
    set_run_font(r, bold=True, size=14)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run(
        "CSDL: winsumwebfinal | MySQL/MariaDB | InnoDB | utf8mb4\n"
        "Phiên bản schema: 1.3 (đã bỏ brands, warehouses)"
    )
    set_run_font(r, size=12)

    doc.add_page_break()

    add_heading(doc, "I. Các bước chuyển thực thể ER thành bảng quan hệ")
    steps = [
        "Bước 1 — Thực thể mạnh → bảng riêng: Mỗi thực thể độc lập có khóa chính riêng. "
        "Ví dụ: thực thể Khách hàng → bảng customers (PK: id).",
        "Bước 2 — Thực thể yếu → bảng phụ thuộc: Bảng con mang FK trỏ về bảng cha. "
        "Ví dụ: Chi tiết đơn hàng → order_items với FK order_id → orders.id.",
        "Bước 3 — Quan hệ 1–N: Thêm FK ở bảng phía «nhiều». "
        "Ví dụ: một Danh mục có nhiều Sản phẩm → products.category_id → categories.id.",
        "Bước 4 — Quan hệ N–N (khi cần): Tạo bảng trung gian. "
        "Trong Winsum Home, giỏ hàng dùng PHP Session thay vì bảng carts.",
        "Bước 5 — Tách thuộc tính đa trị: Ảnh sản phẩm, dòng đơn hàng tách bảng riêng "
        "(product_images, order_items) thay vì lưu JSON trong bảng chính.",
        "Bước 6 — Ràng buộc toàn vẹn: Áp dụng NOT NULL, UNIQUE, ENUM, FK với ON DELETE "
        "CASCADE/SET NULL để đảm bảo dữ liệu nhất quán.",
    ]
    for i, step in enumerate(steps, start=1):
        add_paragraph(doc, f"{i}. {step}")

    add_heading(doc, "II. Lược đồ dữ liệu")
    add_paragraph(
        doc,
        "Lược đồ logic các bảng đang triển khai trong ứng dụng Winsum Home. "
        "Gạch chân = Khóa chính (PK), in nghiêng = Khóa ngoại (FK).",
    )
    add_schema_block(doc)

    add_heading(doc, "III. Mô tả chi tiết cơ sở dữ liệu các bảng")
    add_paragraph(
        doc,
        "Mỗi bảng được mô tả theo mẫu: STT, Tên trường, Kiểu dữ liệu, Độ rộng, Ràng buộc, Mô tả.",
    )

    for tbl in TABLES:
        title = f"Bảng {tbl['no']}: {tbl['name']}"
        add_heading(doc, title, level=2)
        add_paragraph(doc, tbl["purpose"])
        add_schema_table(doc, tbl["fields"])

    doc.add_page_break()

    add_heading(doc, "IV. Lập bảng liên kết (xác định khóa chính khóa phụ)")
    add_paragraph(
        doc,
        "K = Khóa chính (Primary Key) — thuộc tính làm khóa của bảng đó. "
        "C = Khóa liên kết / Khóa ngoại (Foreign Key) — thuộc tính tham chiếu khóa của bảng khác. "
        "Các dòng ghép (order_id, product_id) ghi «dòng loại» vì đã được tách thành từng thuộc tính ở trên.",
    )
    for ent in LINK_TABLE_MAP:
        add_paragraph(doc, f"• {ent}")
    doc.add_paragraph()
    add_link_matrix(doc)
    add_paragraph(
        doc,
        "Giải thích liên kết: customers (1) → orders (4), order_return_requests (7) qua customer_id; "
        "products (2) → order_items (5) qua product_id; categories (3) → products (2) qua category_id; "
        "orders (4) → order_items (5), order_return_requests (7) qua order_id; "
        "coupons (6) → orders (4) qua coupon_id.",
    )

    add_heading(doc, "V. Biểu đồ thực thể – mối quan hệ (ERD)")
    add_paragraph(
        doc,
        "Sơ đồ ERD 7 thực thể cốt lõi của quy trình mua bán Winsum Home "
        "(khách → đặt hàng → chi tiết → hoàn hàng).",
    )
    add_mono_block(doc, ERD_TEXT)

    add_heading(doc, "VI. Lưu trữ ngoài CSDL")
    add_paragraph(
        doc,
        "Giỏ hàng và mã giảm giá tạm thời lưu trong PHP Session "
        "(cart_items, cart_coupon). Khi checkout, dữ liệu chuyển sang bảng orders và order_items.",
    )

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run(
        "Tài liệu sinh tự động từ docs/database/generate-database-doc-word.py — Winsum Home, 06/2026"
    )
    set_run_font(r, size=11, color=RGBColor(0x66, 0x66, 0x66))

    return doc


def main() -> None:
    doc = build_document()
    doc.save(OUTPUT)
    print(str(OUTPUT))


if __name__ == "__main__":
    main()
