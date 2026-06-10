#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Xuất mô tả CSDL Winsum Home ra file Word (.docx)."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

OUTPUT = Path(__file__).resolve().parent / "Mo-ta-CSDL-Winsum-Home.docx"

TABLES = [
    {
        "group": "A. NHÓM NGƯỜI DÙNG",
        "title": "Bảng 1: customers — Khách hàng",
        "purpose": "Lưu tài khoản khách hàng và quản trị viên (role = admin).",
        "fields": [
            ("id", "BIGINT", "PK, AI", "Mã định danh khách hàng"),
            ("customer_code", "VARCHAR(30)", "UK, NOT NULL", "Mã khách hàng (hệ thống sinh)"),
            ("full_name", "VARCHAR(120)", "NOT NULL", "Họ và tên"),
            ("phone", "VARCHAR(30)", "UK, NOT NULL", "Số điện thoại (đăng nhập, liên hệ)"),
            ("password_hash", "VARCHAR(255)", "—", "Mật khẩu đã mã hóa (bcrypt)"),
            ("role", "ENUM", "NOT NULL", "Vai trò: customer hoặc admin"),
        ],
    },
    {
        "group": "B. NHÓM SẢN PHẨM & DANH MỤC",
        "title": "Bảng 2: categories — Danh mục sản phẩm",
        "purpose": "Phân loại sản phẩm theo nhóm (đèn thả, đèn bàn, kệ trang trí…).",
        "fields": [
            ("id", "BIGINT", "PK, AI", "Mã danh mục"),
            ("name", "VARCHAR(120)", "NOT NULL", "Tên danh mục"),
            ("slug", "VARCHAR(140)", "UK, NOT NULL", "Đường dẫn URL danh mục"),
        ],
    },
    {
        "title": "Bảng 3: brands — Thương hiệu",
        "purpose": "Quản lý thương hiệu sản phẩm.",
        "fields": [
            ("id", "BIGINT", "PK, AI", "Mã thương hiệu"),
            ("name", "VARCHAR(120)", "UK, NOT NULL", "Tên thương hiệu"),
            ("slug", "VARCHAR(140)", "UK, NOT NULL", "Slug URL"),
        ],
    },
    {
        "title": "Bảng 4: products — Sản phẩm",
        "purpose": "Thông tin sản phẩm bán trên website.",
        "fields": [
            ("id", "BIGINT", "PK, AI", "Mã sản phẩm"),
            ("category_id", "BIGINT", "FK, NOT NULL", "Danh mục sản phẩm"),
            ("name", "VARCHAR(255)", "NOT NULL", "Tên sản phẩm"),
            ("slug", "VARCHAR(300)", "UK, NOT NULL", "Slug URL sản phẩm"),
            ("sku", "VARCHAR(60)", "UK, NOT NULL", "Mã SKU"),
            ("base_price", "DECIMAL(12,2)", "NOT NULL", "Giá bán"),
        ],
    },
    {
        "title": "Bảng 5: product_images — Ảnh sản phẩm",
        "purpose": "Lưu nhiều ảnh cho một sản phẩm.",
        "fields": [
            ("id", "BIGINT", "PK, AI", "Mã ảnh"),
            ("product_id", "BIGINT", "FK, NOT NULL", "Sản phẩm liên kết"),
            ("image_url", "VARCHAR(255)", "NOT NULL", "Đường dẫn file ảnh"),
        ],
    },
    {
        "group": "C. NHÓM TỒN KHO",
        "title": "Bảng 6: warehouses — Kho hàng",
        "purpose": "Danh sách kho lưu trữ hàng.",
        "fields": [
            ("id", "BIGINT", "PK, AI", "Mã kho"),
            ("name", "VARCHAR(120)", "NOT NULL", "Tên kho"),
            ("code", "VARCHAR(30)", "UK, NOT NULL", "Mã kho (VD: MAIN)"),
        ],
    },
    {
        "title": "Bảng 7: inventory_items — Tồn kho",
        "purpose": "Số lượng tồn theo sản phẩm và kho.",
        "fields": [
            ("id", "BIGINT", "PK, AI", "Mã bản ghi tồn"),
            ("product_id", "BIGINT", "FK, NOT NULL", "Sản phẩm"),
            ("warehouse_id", "BIGINT", "FK, NOT NULL", "Kho lưu trữ"),
            ("quantity_on_hand", "INT", "NOT NULL", "Số lượng tồn thực tế"),
        ],
    },
    {
        "title": "Bảng 8: inventory_alerts — Cảnh báo tồn kho",
        "purpose": "Ghi nhận cảnh báo khi sản phẩm hết hàng.",
        "fields": [
            ("id", "BIGINT", "PK, AI", "Mã cảnh báo"),
            ("product_id", "BIGINT", "FK, NOT NULL", "Sản phẩm hết hàng"),
            ("message", "TEXT", "NOT NULL", "Nội dung cảnh báo"),
        ],
    },
    {
        "group": "D. NHÓM KHUYẾN MÃI",
        "title": "Bảng 9: coupons — Mã giảm giá",
        "purpose": "Quản lý mã giảm giá, freeship, VIP.",
        "fields": [
            ("id", "BIGINT", "PK, AI", "Mã coupon"),
            ("code", "VARCHAR(50)", "UK, NOT NULL", "Mã nhập (VD: WINSUM10)"),
            ("name", "VARCHAR(120)", "NOT NULL", "Tên hiển thị"),
            ("discount_type", "ENUM", "NOT NULL", "fixed / percent / shipping"),
            ("coupon_role", "ENUM", "NOT NULL", "discount / shipping / vip"),
            ("discount_value", "DECIMAL(12,2)", "NOT NULL", "Giá trị giảm (VNĐ hoặc %)"),
        ],
    },
    {
        "title": "Bảng 10: coupon_redemptions — Lượt dùng mã",
        "purpose": "Theo dõi số lần khách sử dụng mã giảm giá.",
        "fields": [
            ("id", "BIGINT", "PK, AI", "Mã bản ghi"),
            ("coupon_id", "BIGINT", "FK, NOT NULL", "Mã giảm giá đã dùng"),
            ("order_id", "BIGINT", "—", "Đơn hàng áp dụng mã"),
        ],
    },
    {
        "group": "E. NHÓM ĐƠN HÀNG",
        "title": "Bảng 11: orders — Đơn hàng",
        "purpose": "Header đơn hàng; lưu snapshot thông tin khách tại thời điểm đặt.",
        "fields": [
            ("id", "BIGINT", "PK, AI", "Mã đơn nội bộ"),
            ("order_code", "VARCHAR(30)", "UK, NOT NULL", "Mã đơn hiển thị"),
            ("customer_name", "VARCHAR(120)", "NOT NULL", "Tên người nhận"),
            ("customer_phone", "VARCHAR(30)", "NOT NULL", "Số điện thoại"),
            ("customer_address", "TEXT", "NOT NULL", "Địa chỉ giao hàng"),
            ("coupon_id", "BIGINT", "FK", "Mã giảm giá (nếu có)"),
            ("coupon_code", "VARCHAR(50)", "—", "Snapshot mã coupon"),
            ("subtotal", "DECIMAL(12,2)", "NOT NULL", "Tổng tiền hàng"),
            ("shipping_fee", "DECIMAL(12,2)", "NOT NULL", "Phí vận chuyển"),
            ("discount_amount", "DECIMAL(12,2)", "NOT NULL", "Số tiền giảm giá"),
            ("grand_total", "DECIMAL(12,2)", "NOT NULL", "Tổng thanh toán"),
            ("payment_status", "ENUM", "NOT NULL", "unpaid / paid / failed / refunded"),
            ("fulfillment_status", "ENUM", "NOT NULL", "Trạng thái giao hàng"),
            ("status", "VARCHAR(30)", "NOT NULL", "Trạng thái đơn tổng hợp"),
        ],
    },
    {
        "title": "Bảng 12: order_items — Chi tiết đơn hàng",
        "purpose": "Snapshot sản phẩm tại thời điểm mua.",
        "fields": [
            ("id", "BIGINT", "PK, AI", "Mã dòng đơn"),
            ("order_id", "BIGINT", "FK, NOT NULL", "Đơn hàng"),
            ("product_sku", "VARCHAR(60)", "NOT NULL", "Mã SKU tại thời điểm mua"),
            ("product_name", "VARCHAR(255)", "NOT NULL", "Tên SP tại thời điểm mua"),
            ("unit_price", "DECIMAL(12,2)", "NOT NULL", "Đơn giá"),
            ("quantity", "INT", "NOT NULL", "Số lượng"),
            ("line_total", "DECIMAL(12,2)", "NOT NULL", "Thành tiền dòng"),
        ],
    },
    {
        "title": "Bảng 13: order_status_histories — Lịch sử trạng thái đơn",
        "purpose": "Ghi log mỗi lần thay đổi trạng thái đơn hàng.",
        "fields": [
            ("id", "BIGINT", "PK, AI", "Mã bản ghi"),
            ("order_id", "BIGINT", "FK, NOT NULL", "Đơn hàng"),
            ("to_status", "VARCHAR(30)", "NOT NULL", "Trạng thái chuyển đến"),
            ("changed_by", "VARCHAR(120)", "—", "customer / admin / vietqr_demo"),
        ],
    },
    {
        "title": "Bảng 14: order_return_requests — Yêu cầu hoàn hàng",
        "purpose": "Khách gửi yêu cầu hoàn trong 7 ngày sau giao hàng.",
        "fields": [
            ("id", "BIGINT", "PK, AI", "Mã yêu cầu"),
            ("order_id", "BIGINT", "FK, NOT NULL", "Đơn cần hoàn"),
            ("customer_id", "BIGINT", "NOT NULL", "Khách gửi yêu cầu"),
            ("reason", "VARCHAR(80)", "NOT NULL", "Lý do hoàn (mã)"),
            ("description", "TEXT", "NOT NULL", "Mô tả chi tiết"),
            ("status", "ENUM", "NOT NULL", "pending / accepted / completed / rejected…"),
        ],
    },
    {
        "group": "F. NHÓM THANH TOÁN & VẬN CHUYỂN",
        "title": "Bảng 15: payment_methods — Phương thức thanh toán",
        "purpose": "Danh mục PT thanh toán (COD, chuyển khoản).",
        "fields": [
            ("id", "BIGINT", "PK, AI", "Mã phương thức"),
            ("code", "VARCHAR(30)", "UK, NOT NULL", "cod / bank_transfer"),
            ("name", "VARCHAR(120)", "NOT NULL", "Tên hiển thị"),
        ],
    },
    {
        "title": "Bảng 16: order_payments — Thanh toán đơn hàng",
        "purpose": "Giao dịch thanh toán gắn với đơn hàng.",
        "fields": [
            ("id", "BIGINT", "PK, AI", "Mã giao dịch"),
            ("order_id", "BIGINT", "FK, NOT NULL", "Đơn hàng"),
            ("payment_method_id", "BIGINT", "FK", "Phương thức thanh toán"),
            ("amount", "DECIMAL(12,2)", "NOT NULL", "Số tiền thanh toán"),
            ("status", "ENUM", "NOT NULL", "pending / success / failed / refunded"),
        ],
    },
    {
        "title": "Bảng 17: shipping_methods — Phương thức vận chuyển",
        "purpose": "Danh mục PT vận chuyển và phí ship.",
        "fields": [
            ("id", "BIGINT", "PK, AI", "Mã phương thức"),
            ("code", "VARCHAR(30)", "UK, NOT NULL", "express_24h / standard"),
            ("name", "VARCHAR(120)", "NOT NULL", "Tên hiển thị"),
            ("fee", "DECIMAL(12,2)", "NOT NULL", "Phí ship cố định"),
        ],
    },
    {
        "title": "Bảng 18: order_shipments — Vận chuyển đơn hàng",
        "purpose": "Thông tin giao hàng của từng đơn.",
        "fields": [
            ("id", "BIGINT", "PK, AI", "Mã vận chuyển"),
            ("order_id", "BIGINT", "FK, NOT NULL", "Đơn hàng"),
            ("shipping_method_id", "BIGINT", "FK", "Phương thức ship"),
            ("recipient_name", "VARCHAR(120)", "NOT NULL", "Tên người nhận"),
            ("recipient_phone", "VARCHAR(30)", "NOT NULL", "SĐT người nhận"),
            ("recipient_address", "TEXT", "NOT NULL", "Địa chỉ nhận hàng"),
            ("shipping_fee", "DECIMAL(12,2)", "NOT NULL", "Phí vận chuyển"),
        ],
    },
    {
        "group": "G. NHÓM ĐÁNH GIÁ & NỘI DUNG",
        "title": "Bảng 19: product_reviews — Đánh giá sản phẩm",
        "purpose": "Khách đánh giá sau khi đơn delivered.",
        "fields": [
            ("id", "BIGINT", "PK, AI", "Mã đánh giá"),
            ("product_id", "BIGINT", "FK, NOT NULL", "Sản phẩm được đánh giá"),
            ("reviewer_name", "VARCHAR(120)", "NOT NULL", "Tên người đánh giá"),
            ("rating", "TINYINT", "NOT NULL", "Điểm từ 1 đến 5"),
        ],
    },
    {
        "title": "Bảng 20: blog_categories — Danh mục blog",
        "purpose": "Phân loại bài viết blog.",
        "fields": [
            ("id", "BIGINT", "PK, AI", "Mã danh mục"),
            ("name", "VARCHAR(120)", "UK, NOT NULL", "Tên danh mục"),
            ("slug", "VARCHAR(140)", "UK, NOT NULL", "Slug URL"),
        ],
    },
    {
        "title": "Bảng 21: blog_posts — Bài viết blog",
        "purpose": "Nội dung bài viết tin tức / hướng dẫn.",
        "fields": [
            ("id", "BIGINT", "PK, AI", "Mã bài viết"),
            ("slug", "VARCHAR(255)", "UK, NOT NULL", "Slug URL bài viết"),
            ("title", "VARCHAR(255)", "NOT NULL", "Tiêu đề"),
            ("excerpt", "TEXT", "NOT NULL", "Tóm tắt ngắn"),
            ("content", "LONGTEXT", "NOT NULL", "Nội dung bài viết"),
            ("image", "VARCHAR(255)", "NOT NULL", "Ảnh đại diện"),
            ("published_at", "DATE", "NOT NULL", "Ngày đăng"),
        ],
    },
    {
        "title": "Bảng 22: blog_comments — Bình luận blog",
        "purpose": "Bình luận bài viết; admin duyệt trước khi hiển thị.",
        "fields": [
            ("id", "BIGINT", "PK, AI", "Mã bình luận"),
            ("post_id", "BIGINT", "NOT NULL", "Bài viết được bình luận"),
            ("author_name", "VARCHAR(120)", "NOT NULL", "Tên người bình luận"),
            ("content", "TEXT", "NOT NULL", "Nội dung bình luận"),
        ],
    },
    {
        "title": "Bảng 23: banners — Banner trang chủ",
        "purpose": "Slider / banner quảng cáo trang chủ.",
        "fields": [
            ("id", "BIGINT", "PK, AI", "Mã banner"),
            ("title", "VARCHAR(160)", "NOT NULL", "Tiêu đề banner"),
            ("position", "VARCHAR(60)", "NOT NULL", "Vị trí hiển thị trên trang"),
        ],
    },
]

FK_ROWS = [
    ("products", "category_id", "categories", "N SP → 1 danh mục"),
    ("product_images", "product_id", "products", "N ảnh → 1 SP"),
    ("inventory_items", "product_id, warehouse_id", "products, warehouses", "Tồn theo SP + kho"),
    ("inventory_alerts", "product_id", "products", "Cảnh báo theo SP"),
    ("orders", "coupon_id", "coupons", "1 đơn → 0/1 mã"),
    ("order_items", "order_id", "orders", "N dòng → 1 đơn"),
    ("order_payments", "order_id, payment_method_id", "orders, payment_methods", "TT theo đơn"),
    ("order_shipments", "order_id, shipping_method_id", "orders, shipping_methods", "Ship theo đơn"),
    ("order_status_histories", "order_id", "orders", "N log → 1 đơn"),
    ("order_return_requests", "order_id", "orders", "N yêu cầu → 1 đơn"),
    ("coupon_redemptions", "coupon_id", "coupons", "N lượt → 1 mã"),
    ("product_reviews", "product_id", "products", "N đánh giá → 1 SP"),
    ("blog_posts", "category_id", "blog_categories", "N bài → 1 DM blog"),
]

SESSION_ROWS = [
    ("Giỏ hàng", "PHP Session $_SESSION['cart_items']", "id SP, tên, SKU, giá, SL, ảnh", "Không có bảng carts"),
    ("Mã giảm giá tạm", "PHP Session $_SESSION['cart_coupon']", "Mã coupon", "Validate qua bảng coupons"),
    ("Sau checkout", "MySQL orders + order_items", "Snapshot giá/tên SP", "Chuyển từ session sang DB"),
]


def set_cell_font(cell, bold=False, size=11):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            run.font.size = Pt(size)
            run.bold = bold
        if not paragraph.runs:
            run = paragraph.add_run(paragraph.text)
            paragraph.text = ""
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            run.font.size = Pt(size)
            run.bold = bold


def add_field_table(doc: Document, fields: list[tuple[str, str, str, str]]) -> None:
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["STT", "Tên trường", "Kiểu dữ liệu", "Ràng buộc", "Mô tả"]
    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr[i].text = text
        set_cell_font(hdr[i], bold=True)
    for idx, (name, dtype, constraint, desc) in enumerate(fields, start=1):
        row = table.add_row().cells
        row[0].text = str(idx)
        row[1].text = name
        row[2].text = dtype
        row[3].text = constraint
        row[4].text = desc
        for cell in row:
            set_cell_font(cell)
    doc.add_paragraph()


def add_simple_table(doc: Document, headers: list[str], rows: list[tuple]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr[i].text = text
        set_cell_font(hdr[i], bold=True)
    for row_data in rows:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = str(val)
            set_cell_font(row[i])
    doc.add_paragraph()


def build_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(13)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("MÔ TẢ CƠ SỞ DỮ LIỆU")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "Times New Roman"

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Website thương mại điện tử Winsum Home")
    r.font.size = Pt(13)
    r.font.name = "Times New Roman"

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run(
        "CSDL: winsumwebfinal | MySQL/MariaDB | InnoDB | utf8mb4 | 23 bảng\n"
        "Nguồn schema: winsumwebfinal (8).sql"
    )
    r.font.size = Pt(12)
    r.font.name = "Times New Roman"

    doc.add_paragraph()

    intro = doc.add_paragraph(
        "Tài liệu mô tả các trường bắt buộc và cốt lõi của từng bảng trong hệ thống. "
        "Các trường tùy chọn (NULL), trường tự sinh (created_at, updated_at) không liệt kê để tránh dài dòng."
    )
    intro.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    current_group = None
    for item in TABLES:
        group = item.get("group")
        if group and group != current_group:
            current_group = group
            h = doc.add_heading(group, level=1)
            for run in h.runs:
                run.font.name = "Times New Roman"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

        h4 = doc.add_heading(item["title"], level=2)
        for run in h4.runs:
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

        p = doc.add_paragraph(item["purpose"])
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_field_table(doc, item["fields"])

    h = doc.add_heading("H. LƯU TRỮ NGOÀI CSDL (GIỎ HÀNG)", level=1)
    for run in h.runs:
        run.font.name = "Times New Roman"

    add_simple_table(
        doc,
        ["Thành phần", "Kiểu lưu trữ", "Dữ liệu bắt buộc", "Ghi chú"],
        SESSION_ROWS,
    )

    h = doc.add_heading("I. BẢNG TÓM TẮT QUAN HỆ KHÓA NGOẠI", level=1)
    for run in h.runs:
        run.font.name = "Times New Roman"

    add_simple_table(
        doc,
        ["Bảng con", "Trường FK", "Bảng cha", "Quan hệ"],
        FK_ROWS,
    )

    footer = doc.add_paragraph(
        "Tài liệu được tạo tự động từ docs/generate-database-doc-word.py — Winsum Home, 2026."
    )
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        run.font.size = Pt(11)
        run.font.name = "Times New Roman"

    return doc


def main() -> None:
    doc = build_document()
    doc.save(OUTPUT)
    print(str(OUTPUT))


if __name__ == "__main__":
    main()
